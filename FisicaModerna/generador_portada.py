#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de Portada para Manual de Prácticas de Laboratorio
Física Moderna - Ingeniería en Nanotecnología
Universidad Tecnológica de Querétaro

Autor: Sistema de Generación Automática de Documentos
Fecha: 2025
Versión: 1.0
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime

class GeneradorPortadaManual:
    """
    Clase para generar la portada del Manual de Prácticas de Laboratorio
    en formato Word (.docx)
    """
    
    def __init__(self):
        """Inicializar el generador con configuraciones por defecto"""
        self.documento = Document()
        
        # Colores institucionales UTEQ
        self.color_azul_uteq = RGBColor(0, 51, 102)      # #003366
        self.color_gris_uteq = RGBColor(128, 128, 128)   # #808080
        self.color_dorado_uteq = RGBColor(255, 204, 0)   # #FFCC00
        self.color_negro = RGBColor(0, 0, 0)             # #000000
        
        # Configuración del documento
        self.configurar_documento()
        
    def configurar_documento(self):
        """Configurar las propiedades generales del documento"""
        # Configurar márgenes
        sections = self.documento.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            
        # Propiedades del documento
        propiedades = self.documento.core_properties
        propiedades.title = "Manual de Prácticas de Laboratorio - Física Moderna"
        propiedades.author = "Universidad Tecnológica de Querétaro"
        propiedades.subject = "Física Moderna - Ingeniería en Nanotecnología"
        propiedades.keywords = "física, cuántica, nanotecnología, prácticas, laboratorio"
        propiedades.created = datetime.now()
        
    def agregar_espaciado(self, lineas=1):
        """Agregar espaciado vertical entre elementos"""
        for _ in range(lineas):
            self.documento.add_paragraph()
            
    def agregar_logo_placeholder(self, texto_placeholder="[LOGO UTEQ]"):
        """Agregar placeholder para el logo institucional"""
        logo_para = self.documento.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = logo_para.add_run(texto_placeholder)
        run.font.size = Pt(14)
        run.font.color.rgb = self.color_gris_uteq
        run.italic = True
        
        # Nota: Para insertar imagen real, descomenta y ajusta la siguiente línea:
        logo_para.add_run().add_picture('C:\\Repositorios\\CursosUTEQ\\Imagenes\\Logo_uteq.png', width=Inches(1.5))

    def agregar_titulo_principal(self):
        """Agregar el título principal del manual"""
        # Título "MANUAL DE PRÁCTICAS"
        titulo1 = self.documento.add_paragraph()
        titulo1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run1 = titulo1.add_run("MANUAL DE PRÁCTICAS")
        run1.font.name = "Arial"
        run1.font.size = Pt(28)
        run1.font.color.rgb = self.color_azul_uteq
        run1.font.bold = True
        
        # Título "DE LABORATORIO"
        titulo2 = self.documento.add_paragraph()
        titulo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run2 = titulo2.add_run("DE LABORATORIO")
        run2.font.name = "Arial"
        run2.font.size = Pt(28)
        run2.font.color.rgb = self.color_azul_uteq
        run2.font.bold = True
        
        self.agregar_espaciado(1)
        
        # Subtítulo "FÍSICA MODERNA"
        subtitulo = self.documento.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run3 = subtitulo.add_run("FÍSICA MODERNA")
        run3.font.name = "Arial"
        run3.font.size = Pt(20)
        run3.font.color.rgb = self.color_gris_uteq
        run3.font.bold = True
        
        # Descripción
        descripcion = self.documento.add_paragraph()
        descripcion.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run4 = descripcion.add_run("Fundamentos de Teoría Cuántica y Aplicaciones en Nanotecnología")
        run4.font.name = "Arial"
        run4.font.size = Pt(14)
        run4.font.color.rgb = self.color_gris_uteq
        run4.italic = True
        
    def agregar_informacion_institucional(self):
        """Agregar información de la universidad y carrera"""
        self.agregar_espaciado(1)
        
        # Universidad
        uni_para = self.documento.add_paragraph()
        uni_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run1 = uni_para.add_run("UNIVERSIDAD TECNOLÓGICA DE QUERÉTARO")
        run1.font.name = "Arial"
        run1.font.size = Pt(16)
        run1.font.color.rgb = self.color_azul_uteq
        run1.font.bold = True
        
        # División
        div_para = self.documento.add_paragraph()
        div_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run2 = div_para.add_run("División Académica de Tecnologías de la Información y Comunicación")
        run2.font.name = "Arial"
        run2.font.size = Pt(12)
        run2.font.color.rgb = self.color_gris_uteq
        
        # Carrera
        carrera_para = self.documento.add_paragraph()
        carrera_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run3 = carrera_para.add_run("Ingeniería en Nanotecnología")
        run3.font.name = "Arial"
        run3.font.size = Pt(12)
        run3.font.color.rgb = self.color_gris_uteq
        
    def agregar_tabla_informacion_curso(self):
        """Agregar tabla con información del curso"""
        self.agregar_espaciado(2)
        
        # Crear tabla
        tabla = self.documento.add_table(rows=7, cols=2)
        tabla.style = 'Table Grid'
        
        # Datos del curso
        datos_curso = [
            ("Clave:", "FIS-902"),
            ("Créditos:", "5"),
            ("Horas teóricas:", "24"),
            ("Horas prácticas:", "36"),
            ("Horas totales:", "60"),
            ("Cuatrimestre:", "Noveno"),
            ("Modalidad:", "Presencial asistida por tecnología")
        ]
        
        # Llenar tabla
        for i, (campo, valor) in enumerate(datos_curso):
            # Celda del campo
            celda_campo = tabla.cell(i, 0)
            celda_campo.text = campo
            for paragraph in celda_campo.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = self.color_azul_uteq
                    
            # Celda del valor
            celda_valor = tabla.cell(i, 1)
            celda_valor.text = valor
            for paragraph in celda_valor.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.color.rgb = self.color_negro
        
        # Centrar tabla
        tabla.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def agregar_informacion_cuatrimestre(self):
        """Agregar información del cuatrimestre y fechas"""
        self.agregar_espaciado(2)
        
        # Cuatrimestre
        cuatri_para = self.documento.add_paragraph()
        cuatri_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run1 = cuatri_para.add_run("Cuatrimestre Mayo - Agosto 2025")
        run1.font.name = "Arial"
        run1.font.size = Pt(16)
        run1.font.color.rgb = self.color_gris_uteq
        run1.font.bold = True
        
        self.agregar_espaciado(1)
        
        # Lineamientos
        lineamiento1 = self.documento.add_paragraph()
        lineamiento1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run2 = lineamiento1.add_run("Elaborado conforme a los lineamientos de la Nueva Escuela Mexicana")
        run2.font.name = "Arial"
        run2.font.size = Pt(10)
        run2.font.color.rgb = self.color_gris_uteq
        
        lineamiento2 = self.documento.add_paragraph()
        lineamiento2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run3 = lineamiento2.add_run("y las recomendaciones de la UNESCO sobre educación técnica")
        run3.font.name = "Arial"
        run3.font.size = Pt(10)
        run3.font.color.rgb = self.color_gris_uteq
        
    def agregar_logos_institucionales_footer(self):
        """Agregar placeholders para logos institucionales en el pie"""
        self.agregar_espaciado(2)
        
        logos_para = self.documento.add_paragraph()
        logos_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Placeholder para logos
        run = logos_para.add_run("[LOGO SEP]    [LOGO UTP]    [LOGO UTEQ]")
        run.font.size = Pt(12)
        run.font.color.rgb = self.color_gris_uteq
        run.italic = True
        
        # Nota: Para insertar imágenes reales, usar:
        # logos_para.add_run().add_picture('logo_sep.png', width=Inches(0.8))
        # logos_para.add_run("    ")
        # logos_para.add_run().add_picture('logo_utp.png', width=Inches(0.8))
        # logos_para.add_run("    ")
        # logos_para.add_run().add_picture('logo_uteq.png', width=Inches(0.8))
        
    def generar_portada_completa(self, nombre_archivo="Manual_Practicas_Fisica_Moderna_Portada.docx"):
        """
        Generar la portada completa del manual
        
        Args:
            nombre_archivo (str): Nombre del archivo a generar
            
        Returns:
            str: Ruta del archivo generado
        """
        print("🔄 Generando portada del Manual de Prácticas...")
        
        # Logo superior
        self.agregar_espaciado(2)
        self.agregar_logo_placeholder()
        
        # Información institucional superior
        self.agregar_espaciado(1)
        self.agregar_informacion_institucional()
        
        # Título principal
        self.agregar_espaciado(3)
        self.agregar_titulo_principal()
        
        # Tabla de información del curso
        self.agregar_tabla_informacion_curso()
        
        # Información del cuatrimestre
        self.agregar_informacion_cuatrimestre()
        
        # Logos institucionales footer
        self.agregar_logos_institucionales_footer()
        
        # Guardar documento
        ruta_completa = os.path.abspath(nombre_archivo)
        self.documento.save(nombre_archivo)
        
        print(f"✅ Portada generada exitosamente: {ruta_completa}")
        return ruta_completa
        
    def personalizar_portada(self, datos_personalizados):
        """
        Personalizar la portada con datos específicos
        
        Args:
            datos_personalizados (dict): Diccionario con los datos a personalizar
                - universidad: str
                - division: str  
                - carrera: str
                - asignatura: str
                - descripcion: str
                - cuatrimestre: str
                - autor: str
                - etc.
        """
        # Esta función permite personalizar todos los elementos
        # Implementar según necesidades específicas
        pass

def main():
    """Función principal para generar la portada"""
    print("📚 Generador de Portada - Manual de Prácticas de Laboratorio")
    print("=" * 60)
    
    try:
        # Crear instancia del generador
        generador = GeneradorPortadaManual()
        
        # Generar portada
        archivo_generado = generador.generar_portada_completa()
        
        print("\n📋 Información del archivo generado:")
        print(f"   📁 Ubicación: {archivo_generado}")
        print(f"   📄 Formato: Microsoft Word (.docx)")
        print(f"   🎨 Incluye: Colores institucionales, logos placeholder, tabla de información")
        
        print("\n🔧 Para personalizar:")
        print("   1. Reemplazar placeholders de logos con imágenes reales")
        print("   2. Ajustar colores institucionales si es necesario")
        print("   3. Modificar datos del curso según corresponda")
        print("   4. Agregar información adicional requerida")
        
        print("\n✨ ¡Portada generada exitosamente!")
        
    except Exception as e:
        print(f"❌ Error al generar la portada: {e}")
        print("💡 Asegúrate de tener instalada la librería python-docx:")
        print("   pip install python-docx")

def generar_version_personalizada():
    """Generar una versión personalizada de la portada"""
    
    # Datos personalizados (ejemplo)
    datos_custom = {
        "universidad": "Universidad Tecnológica de Querétaro",
        "division": "División Académica de Tecnologías de la Información y Comunicación",
        "carrera": "Ingeniería en Nanotecnología",
        "asignatura": "Física Moderna",
        "descripcion": "Fundamentos de Teoría Cuántica y Aplicaciones en Nanotecnología",
        "cuatrimestre": "Mayo - Agosto 2025",
        "autor": "Dr. [Nombre del Professor]",
        "version": "Primera Edición"
    }
    
    generador = GeneradorPortadaManual()
    archivo = generador.generar_portada_completa("Manual_Personalizado_Portada.docx")
    return archivo

def agregar_pagina_creditos(documento):
    """
    Agregar página de créditos después de la portada
    
    Args:
        documento: Objeto Document de python-docx
    """
    # Agregar salto de página
    documento.add_page_break()
    
    # Título de créditos
    titulo_creditos = documento.add_paragraph()
    titulo_creditos.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run_titulo = titulo_creditos.add_run("CRÉDITOS Y RECONOCIMIENTOS")
    run_titulo.font.name = "Arial"
    run_titulo.font.size = Pt(18)
    run_titulo.font.bold = True
    run_titulo.font.color.rgb = RGBColor(0, 51, 102)
    
    # Agregar contenido de créditos
    # (Similar estructura a la portada)
    
if __name__ == "__main__":
    main()
    
    # Opcional: generar versión personalizada
    # generar_version_personalizada()